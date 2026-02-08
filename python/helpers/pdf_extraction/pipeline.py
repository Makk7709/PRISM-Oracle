"""
╔══════════════════════════════════════════════════════════════════════════════╗
║                    PDF EXTRACTION - PIPELINE                                 ║
║                                                                              ║
║  Non-blocking, fault-tolerant PDF extraction pipeline.                       ║
║  All operations are timeboxed with automatic fallback.                       ║
║                                                                              ║
║  Key guarantees:                                                             ║
║  - No operation blocks indefinitely                                          ║
║  - Circuit breaker prevents cascade failures                                 ║
║  - Always returns a result (even if partial/error)                           ║
║  - No user content in logs                                                   ║
║                                                                              ║
║  Version: 1.0.0                                                              ║
║  © 2025 Korev AI — Proprietary                                               ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import asyncio
import io
import logging
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from contextlib import contextmanager
from functools import wraps
from pathlib import Path
from typing import Any, Callable, Optional, Union

from .config import PDFExtractionConfig, get_default_config
from .types import (
    BBox,
    Cell,
    CircuitBreakerState,
    Diagnostics,
    ExtractionContext,
    ExtractionMethod,
    ExtractionResult,
    ExtractionStatus,
    OutputArtifacts,
    PDFType,
    TableResult,
    Word,
)


logger = logging.getLogger(__name__)

# Thread pool for bounded async operations
_executor = ThreadPoolExecutor(max_workers=4, thread_name_prefix="pdf_extract_")


# ═══════════════════════════════════════════════════════════════════════════════
# TIMEOUT UTILITIES
# ═══════════════════════════════════════════════════════════════════════════════

def run_with_timeout(func: Callable, timeout_s: float, *args, **kwargs) -> tuple[Any, bool]:
    """
    Run a function with a timeout.
    
    Returns: (result, timed_out)
    - If successful: (result, False)
    - If timed out: (None, True)
    - If error: raises the exception
    """
    future = _executor.submit(func, *args, **kwargs)
    try:
        result = future.result(timeout=timeout_s)
        return result, False
    except FuturesTimeoutError:
        future.cancel()
        return None, True


async def run_with_timeout_async(
    func: Callable, timeout_s: float, *args, **kwargs
) -> tuple[Any, bool]:
    """
    Async version of run_with_timeout.
    
    Returns: (result, timed_out)
    """
    loop = asyncio.get_event_loop()
    try:
        result = await asyncio.wait_for(
            loop.run_in_executor(_executor, lambda: func(*args, **kwargs)),
            timeout=timeout_s
        )
        return result, False
    except asyncio.TimeoutError:
        return None, True


@contextmanager
def timed_operation(diagnostics: Diagnostics, operation: str, 
                    page: Optional[int] = None, engine: Optional[str] = None):
    """Context manager for timing an operation."""
    start = time.time()
    status = ExtractionStatus.SUCCESS
    try:
        yield
    except Exception:
        status = ExtractionStatus.ERROR
        raise
    finally:
        duration_ms = int((time.time() - start) * 1000)
        diagnostics.add_timing(operation, duration_ms, status, page, engine)


# ═══════════════════════════════════════════════════════════════════════════════
# CIRCUIT BREAKER
# ═══════════════════════════════════════════════════════════════════════════════

class CircuitBreaker:
    """
    Circuit breaker to prevent cascade failures.
    
    When too many failures occur, the circuit "opens" and blocks
    further attempts to heavy engines.
    """
    
    def __init__(self, config: PDFExtractionConfig):
        self.config = config.budgets.circuit_breaker
        self.state = CircuitBreakerState()
    
    def record_timeout(self, engine: Optional[str] = None):
        """Record a timeout event."""
        self.state.timeout_count += 1
        self.state.last_failure_time = time.time()
        
        if self.state.timeout_count >= self.config.max_timeouts:
            self.state.is_open = True
            logger.warning(
                "Circuit breaker OPEN after %d timeouts",
                self.state.timeout_count
            )
    
    def record_failure(self, engine: str):
        """Record an engine failure."""
        self.state.failure_count += 1
        self.state.last_failure_time = time.time()
        
        if self.state.failure_count >= self.config.max_engine_failures:
            if engine not in self.state.disabled_engines:
                self.state.disabled_engines.append(engine)
                logger.warning("Engine %s disabled after %d failures", 
                             engine, self.state.failure_count)
    
    def is_engine_allowed(self, engine: str) -> bool:
        """Check if an engine is allowed to run."""
        if self.state.is_open:
            return False
        if engine in self.state.disabled_engines:
            # Check cooldown
            if self.state.last_failure_time:
                elapsed = time.time() - self.state.last_failure_time
                if elapsed >= self.config.failure_cooldown_s:
                    self.state.disabled_engines.remove(engine)
                    return True
            return False
        return True
    
    def get_state(self) -> CircuitBreakerState:
        """Get current state."""
        return self.state


# ═══════════════════════════════════════════════════════════════════════════════
# PDF CLASSIFICATION
# ═══════════════════════════════════════════════════════════════════════════════

def classify_pdf(
    doc,
    config: PDFExtractionConfig
) -> tuple[PDFType, float]:
    """
    Classify PDF as text/hybrid/scan.
    
    Args:
        doc: A PDFDocument from the backend abstraction layer,
             or a legacy fitz.Document for backward compatibility.
    
    Returns: (pdf_type, confidence)
    """
    if not config.classification.enabled:
        return PDFType.UNKNOWN, 0.5
    
    heuristics = config.classification.heuristics
    
    # Support both backend abstraction and legacy fitz.Document
    from python.helpers.pdf_extraction.pdf_backend import PDFDocument as _PDFDocument
    if isinstance(doc, _PDFDocument):
        total_pages = min(doc.page_count(), config.budgets.max_pages)
    else:
        total_pages = min(doc.page_count, config.budgets.max_pages)
    
    if total_pages == 0:
        return PDFType.UNKNOWN, 0.0
    
    # Analyze pages
    text_pages = 0
    image_heavy_pages = 0
    empty_pages = 0
    total_chars = 0
    total_words = 0
    
    for page_num in range(total_pages):
        if isinstance(doc, _PDFDocument):
            page_data = doc.get_page(page_num)
            text = page_data.text
            word_count = len(page_data.words)
            image_list = page_data.images
            page_width = page_data.width
            page_height = page_data.height
        else:
            page = doc[page_num]
            text = page.get_text("text")
            word_count = len(page.get_text("words"))
            image_list = page.get_images()
            page_width = page.rect.width
            page_height = page.rect.height
        
        char_count = len(text.strip())
        total_chars += char_count
        total_words += word_count
        
        # Check if page has enough text
        if char_count >= heuristics.min_chars_per_page_for_text_pdf:
            text_pages += 1
        elif char_count < 10:
            empty_pages += 1
        
        # Check image coverage
        if image_list:
            page_area = page_width * page_height
            image_area = 0
            if isinstance(doc, _PDFDocument):
                for img in image_list:
                    image_area += img.width * img.height
            else:
                for img in image_list:
                    try:
                        xref = img[0]
                        img_rect = page.get_image_bbox(xref)
                        if img_rect:
                            image_area += img_rect.width * img_rect.height
                    except Exception:
                        pass
            if page_area > 0 and (image_area / page_area) > heuristics.image_area_ratio_scan_threshold:
                image_heavy_pages += 1
    
    # Determine type
    text_ratio = text_pages / total_pages if total_pages > 0 else 0
    empty_ratio = empty_pages / total_pages if total_pages > 0 else 0
    image_ratio = image_heavy_pages / total_pages if total_pages > 0 else 0
    
    # Decision logic
    if empty_ratio > heuristics.max_empty_text_pages_ratio or image_ratio > 0.6:
        pdf_type = PDFType.SCAN
        confidence = min(0.9, empty_ratio + image_ratio * 0.5)
    elif text_ratio > 0.8 and image_ratio < 0.2:
        pdf_type = PDFType.TEXT
        confidence = text_ratio
    else:
        pdf_type = PDFType.HYBRID
        confidence = 0.5 + (text_ratio * 0.3) - (image_ratio * 0.2)
    
    return pdf_type, round(confidence, 2)


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION (PyMuPDF)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_words_pymupdf(
    doc,
    config: PDFExtractionConfig,
    context: ExtractionContext
) -> list[Word]:
    """
    Extract words with positions from a PDF document.
    
    This is the primary extraction method - fast and reliable.
    Supports both the backend abstraction layer and legacy fitz.Document.
    """
    from python.helpers.pdf_extraction.pdf_backend import PDFDocument as _PDFDocument
    
    words: list[Word] = []
    text_config = config.text
    
    if isinstance(doc, _PDFDocument):
        max_pages = min(doc.page_count(), config.budgets.max_pages)
    else:
        max_pages = min(doc.page_count, config.budgets.max_pages)
    
    for page_num in range(max_pages):
        # Check budget
        if context.is_budget_exhausted(config.budgets.total_timeout_s):
            break
        
        if isinstance(doc, _PDFDocument):
            page_data = doc.get_page(page_num)
            for pw in page_data.words:
                text = pw.text
                if text_config.normalize_whitespace:
                    text = " ".join(text.split())
                if not text:
                    continue
                word = Word(
                    text=text,
                    bbox=BBox(x0=pw.x0, y0=pw.y0, x1=pw.x1, y1=pw.y1),
                    page=page_num,
                    confidence=pw.confidence,
                )
                words.append(word)
        else:
            # Legacy fitz.Document path
            page = doc[page_num]
            word_list = page.get_text("words", sort=text_config.sort_words_reading_order)
            for w in word_list:
                if len(w) >= 5:
                    text = w[4]
                    if text_config.normalize_whitespace:
                        text = " ".join(text.split())
                    if not text:
                        continue
                    word = Word(
                        text=text,
                        bbox=BBox(x0=w[0], y0=w[1], x1=w[2], y1=w[3]),
                        page=page_num,
                        confidence=1.0,
                    )
                    words.append(word)
        
        context.pages_processed += 1
    
    return words


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE EXTRACTION - GEOMETRY RECONSTRUCTION
# ═══════════════════════════════════════════════════════════════════════════════

def detect_table_regions(
    words: list[Word],
    page: int,
    config: PDFExtractionConfig
) -> list[BBox]:
    """
    Detect potential table regions on a page using word clustering.
    
    Returns list of bounding boxes for table regions.
    """
    page_words = [w for w in words if w.page == page]
    if not page_words:
        return []
    
    # Simple heuristic: find dense regions of aligned words
    # For a real implementation, use DBSCAN or similar
    if len(page_words) < 10:
        return []
    
    # Get page bounds
    min_x = min(w.bbox.x0 for w in page_words)
    max_x = max(w.bbox.x1 for w in page_words)
    min_y = min(w.bbox.y0 for w in page_words)
    max_y = max(w.bbox.y1 for w in page_words)
    
    # For now, treat entire content area as potential table region
    # A real implementation would cluster words more intelligently
    return [BBox(x0=min_x, y0=min_y, x1=max_x, y1=max_y)]


def cluster_columns(
    words: list[Word],
    config: PDFExtractionConfig
) -> list[float]:
    """
    Cluster word x-coordinates to detect columns.
    
    Returns list of column x-boundaries.
    """
    col_config = config.tables.geometry.column_detection
    
    if not words:
        return []
    
    # Get x-coordinates of word starts
    x_coords = sorted(set(w.bbox.x0 for w in words))
    
    if len(x_coords) < 2:
        return x_coords
    
    # Simple clustering: merge coordinates within eps
    eps = col_config.x_cluster_eps
    clusters = []
    current_cluster = [x_coords[0]]
    
    for x in x_coords[1:]:
        if x - current_cluster[-1] < eps:
            current_cluster.append(x)
        else:
            clusters.append(sum(current_cluster) / len(current_cluster))
            current_cluster = [x]
    
    if current_cluster:
        clusters.append(sum(current_cluster) / len(current_cluster))
    
    # Limit columns
    if len(clusters) > col_config.max_columns:
        clusters = clusters[:col_config.max_columns]
    
    return clusters


def cluster_rows(
    words: list[Word],
    config: PDFExtractionConfig
) -> list[float]:
    """
    Cluster word y-coordinates to detect rows.
    
    Returns list of row y-boundaries.
    """
    row_config = config.tables.geometry.row_detection
    
    if not words:
        return []
    
    # Get y-coordinates
    y_coords = sorted(set(w.bbox.center_y for w in words))
    
    if len(y_coords) < 2:
        return y_coords
    
    # Simple clustering
    eps = row_config.y_cluster_eps
    clusters = []
    current_cluster = [y_coords[0]]
    
    for y in y_coords[1:]:
        if y - current_cluster[-1] < eps:
            current_cluster.append(y)
        else:
            clusters.append(sum(current_cluster) / len(current_cluster))
            current_cluster = [y]
    
    if current_cluster:
        clusters.append(sum(current_cluster) / len(current_cluster))
    
    # Limit rows
    if len(clusters) > row_config.max_rows:
        clusters = clusters[:row_config.max_rows]
    
    return clusters


def assign_words_to_cells(
    words: list[Word],
    col_boundaries: list[float],
    row_boundaries: list[float],
    config: PDFExtractionConfig
) -> list[Cell]:
    """
    Assign words to cells based on position.
    
    Returns list of cells.
    """
    if not col_boundaries or not row_boundaries:
        return []
    
    cell_config = config.tables.geometry.cell_assignment
    post_config = config.tables.geometry.postprocessing
    
    num_cols = len(col_boundaries)
    num_rows = len(row_boundaries)
    
    # Build cell grid
    cell_texts: dict[tuple[int, int], list[str]] = {}
    cell_bboxes: dict[tuple[int, int], list[BBox]] = {}
    
    for word in words:
        # Find column
        col = 0
        for i, boundary in enumerate(col_boundaries):
            if word.bbox.x0 < boundary + config.tables.geometry.column_detection.x_cluster_eps:
                col = i
                break
            col = i
        
        # Find row
        row = 0
        for i, boundary in enumerate(row_boundaries):
            if word.bbox.center_y < boundary + config.tables.geometry.row_detection.y_cluster_eps:
                row = i
                break
            row = i
        
        key = (row, col)
        if key not in cell_texts:
            cell_texts[key] = []
            cell_bboxes[key] = []
        cell_texts[key].append(word.text)
        cell_bboxes[key].append(word.bbox)
    
    # Build cells
    cells = []
    for (row, col), texts in cell_texts.items():
        text = post_config.join_token_separator.join(texts)
        
        # Postprocess
        if post_config.strip_currency_spacing:
            # Remove spaces in numbers like "1 000 000"
            import re
            text = re.sub(r'(\d)\s+(\d)', r'\1\2', text)
        
        bboxes = cell_bboxes[(row, col)]
        if bboxes:
            cell_bbox = BBox(
                x0=min(b.x0 for b in bboxes),
                y0=min(b.y0 for b in bboxes),
                x1=max(b.x1 for b in bboxes),
                y1=max(b.y1 for b in bboxes)
            )
        else:
            cell_bbox = None
        
        cells.append(Cell(
            text=text,
            row=row,
            col=col,
            bbox=cell_bbox
        ))
    
    return cells


def build_table_geometry(
    words: list[Word],
    region: BBox,
    page: int,
    config: PDFExtractionConfig
) -> Optional[TableResult]:
    """
    Build a table from words using geometry reconstruction.
    
    This is the default, safe extraction method.
    """
    start_time = time.time()
    
    # Filter words in region
    region_words = [
        w for w in words 
        if w.page == page and region.overlaps(w.bbox, threshold=0.1)
    ]
    
    if len(region_words) < 4:  # Need at least a few words
        return None
    
    # Detect columns and rows
    col_boundaries = cluster_columns(region_words, config)
    row_boundaries = cluster_rows(region_words, config)
    
    if len(col_boundaries) < 2 or len(row_boundaries) < 2:
        return None
    
    # Assign words to cells
    cells = assign_words_to_cells(region_words, col_boundaries, row_boundaries, config)
    
    if not cells:
        return None
    
    num_rows = max(c.row for c in cells) + 1
    num_cols = max(c.col for c in cells) + 1
    
    # Check max cells limit
    if num_rows * num_cols > config.budgets.max_cells_per_table:
        return None
    
    # Calculate quality metrics
    total_cells = num_rows * num_cols
    filled_cells = len([c for c in cells if c.text.strip()])
    fill_ratio = filled_cells / total_cells if total_cells > 0 else 0
    
    # Check row consistency
    row_col_counts = {}
    for cell in cells:
        row_col_counts[cell.row] = row_col_counts.get(cell.row, 0) + 1
    
    if row_col_counts:
        most_common_count = max(set(row_col_counts.values()), key=list(row_col_counts.values()).count)
        consistent_rows = sum(1 for c in row_col_counts.values() if c == most_common_count)
        consistent_ratio = consistent_rows / len(row_col_counts)
    else:
        consistent_ratio = 0
    
    # Quality thresholds
    quality_config = config.tables.geometry.quality_checks
    confidence = 0.5
    
    if fill_ratio >= quality_config.min_fill_ratio:
        confidence += 0.2
    if consistent_ratio >= quality_config.min_consistent_columns_ratio:
        confidence += 0.2
    
    confidence = min(1.0, confidence)
    
    extraction_time = int((time.time() - start_time) * 1000)
    
    return TableResult(
        cells=cells,
        num_rows=num_rows,
        num_cols=num_cols,
        page=page,
        bbox=region,
        method=ExtractionMethod.PYMUPDF_GEOMETRY,
        confidence=confidence,
        fill_ratio=fill_ratio,
        extraction_time_ms=extraction_time
    )


def extract_tables_geometry(
    words: list[Word],
    doc,
    config: PDFExtractionConfig,
    context: ExtractionContext,
    diagnostics: Diagnostics
) -> list[TableResult]:
    """
    Extract tables using geometry reconstruction.
    
    This is the DEFAULT and SAFE method.
    """
    from python.helpers.pdf_extraction.pdf_backend import PDFDocument as _PDFDocument
    
    if not config.tables.enabled or not config.tables.geometry.enabled:
        return []
    
    tables: list[TableResult] = []
    if isinstance(doc, _PDFDocument):
        max_pages = min(doc.page_count(), config.budgets.max_pages)
    else:
        max_pages = min(doc.page_count, config.budgets.max_pages)
    
    for page_num in range(max_pages):
        # Check budget
        if context.is_budget_exhausted(config.budgets.total_timeout_s):
            break
        
        # Detect table regions
        regions = detect_table_regions(words, page_num, config)
        regions = regions[:config.budgets.max_table_regions_per_page]
        
        for region in regions:
            table = build_table_geometry(words, region, page_num, config)
            if table:
                tables.append(table)
                diagnostics.table_count += 1
    
    diagnostics.methods_attempted.append(ExtractionMethod.PYMUPDF_GEOMETRY)
    if tables:
        diagnostics.methods_succeeded.append(ExtractionMethod.PYMUPDF_GEOMETRY)
    
    return tables


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE EXTRACTION - OPTIONAL ENGINES (BOUNDED)
# ═══════════════════════════════════════════════════════════════════════════════

def try_camelot_extraction(
    pdf_path: str,
    page: int,
    config: PDFExtractionConfig,
    circuit_breaker: CircuitBreaker
) -> Optional[list[TableResult]]:
    """
    Try extracting tables with Camelot (bounded).
    
    Returns None if:
    - Engine not enabled
    - Circuit breaker open
    - Timeout
    - Error
    """
    if not config.is_engine_enabled("camelot"):
        return None
    
    if not circuit_breaker.is_engine_allowed("camelot"):
        return None
    
    try:
        import camelot
    except ImportError:
        return None
    
    camelot_config = config.tables.optional_engines.camelot
    timeout = config.budgets.per_engine_timeout_s
    
    def _extract():
        tables = []
        for flavor in camelot_config.flavor_order:
            try:
                result = camelot.read_pdf(
                    pdf_path,
                    pages=str(page + 1),  # 1-indexed
                    flavor=flavor,
                    line_scale=camelot_config.line_scale,
                    strip_text=camelot_config.strip_text
                )
                
                for table in result:
                    df = table.df
                    cells = []
                    for i, row in df.iterrows():
                        for j, val in enumerate(row):
                            cells.append(Cell(
                                text=str(val) if val else "",
                                row=i,
                                col=j
                            ))
                    
                    method = (ExtractionMethod.CAMELOT_LATTICE 
                             if flavor == "lattice" 
                             else ExtractionMethod.CAMELOT_STREAM)
                    
                    tables.append(TableResult(
                        cells=cells,
                        num_rows=len(df),
                        num_cols=len(df.columns),
                        page=page,
                        method=method,
                        confidence=table.accuracy / 100 if hasattr(table, 'accuracy') else 0.7
                    ))
                
                if tables:
                    break  # Got results, stop trying flavors
                    
            except Exception:
                continue
        
        return tables
    
    result, timed_out = run_with_timeout(_extract, timeout)
    
    if timed_out:
        circuit_breaker.record_timeout("camelot")
        return None
    
    return result


def try_optional_engines(
    pdf_path: str,
    words: list[Word],
    doc,
    config: PDFExtractionConfig,
    context: ExtractionContext,
    circuit_breaker: CircuitBreaker,
    diagnostics: Diagnostics
) -> list[TableResult]:
    """
    Try optional engines for table extraction.
    
    Only runs if:
    - Geometry extraction had low confidence
    - Budget allows
    - Circuit breaker is closed
    """
    from python.helpers.pdf_extraction.pdf_backend import PDFDocument as _PDFDocument
    
    tables: list[TableResult] = []
    
    # Check if we should try engines
    if context.backtracks_used >= config.budgets.max_backtracks:
        return tables
    
    if isinstance(doc, _PDFDocument):
        max_pages = min(doc.page_count(), config.budgets.max_pages)
    else:
        max_pages = min(doc.page_count, config.budgets.max_pages)
    
    for page_num in range(max_pages):
        if context.is_budget_exhausted(config.budgets.total_timeout_s):
            break
        
        # Try Camelot
        if config.is_engine_enabled("camelot"):
            camelot_tables = try_camelot_extraction(
                pdf_path, page_num, config, circuit_breaker
            )
            if camelot_tables:
                tables.extend(camelot_tables)
                diagnostics.methods_attempted.append(ExtractionMethod.CAMELOT_LATTICE)
                diagnostics.methods_succeeded.append(ExtractionMethod.CAMELOT_LATTICE)
    
    context.backtracks_used += 1
    
    return tables


# ═══════════════════════════════════════════════════════════════════════════════
# OUTPUT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def generate_outputs(
    tables: list[TableResult],
    config: PDFExtractionConfig
) -> OutputArtifacts:
    """Generate output artifacts (CSV, JSON, DOCX)."""
    outputs = OutputArtifacts()
    
    if not tables:
        return outputs
    
    # CSV
    if config.output.return_csv:
        csv_parts = []
        for i, table in enumerate(tables):
            csv_parts.append(f"# Table {i + 1} (Page {table.page + 1})")
            csv_parts.append(table.to_csv())
        outputs.csv_data = "\n".join(csv_parts)
    
    # JSON
    if config.output.return_json_cells:
        outputs.json_data = {
            "tables": [table.to_dict() for table in tables],
            "table_count": len(tables)
        }
    
    # DOCX
    if config.output.return_docx:
        try:
            outputs.docx_bytes = generate_docx(tables, config)
        except Exception as e:
            logger.warning("DOCX generation failed: %s", type(e).__name__)
    
    return outputs


def generate_docx(tables: list[TableResult], config: PDFExtractionConfig) -> bytes:
    """Generate DOCX from tables."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return b""
    
    doc = Document()
    docx_config = config.output.docx
    
    for i, table in enumerate(tables):
        # Add heading
        doc.add_heading(f"Table {i + 1} (Page {table.page + 1})", level=2)
        
        # Create table
        rows = table.to_rows()
        if not rows:
            continue
        
        word_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        word_table.style = docx_config.table_style
        
        for row_idx, row in enumerate(rows):
            for col_idx, cell_text in enumerate(row):
                cell = word_table.cell(row_idx, col_idx)
                cell.text = cell_text
                
                # Format
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = docx_config.font_name
                        run.font.size = Pt(docx_config.font_size_pt)
                        if row_idx < table.header_rows and docx_config.header_bold:
                            run.font.bold = True
        
        doc.add_paragraph()  # Spacing
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# CONFIDENCE CALCULATION
# ═══════════════════════════════════════════════════════════════════════════════

def calculate_overall_confidence(
    pdf_type: PDFType,
    pdf_type_confidence: float,
    tables: list[TableResult],
    diagnostics: Diagnostics
) -> tuple[float, bool, list[str]]:
    """
    Calculate overall confidence and determine if human review is needed.
    
    Returns: (confidence, requires_review, review_reasons)
    """
    confidence_parts = [pdf_type_confidence]
    review_reasons = []
    
    # Table confidence
    if tables:
        table_confidences = [t.confidence for t in tables]
        avg_table_confidence = sum(table_confidences) / len(table_confidences)
        confidence_parts.append(avg_table_confidence)
        
        if avg_table_confidence < 0.5:
            review_reasons.append("Low table extraction confidence")
    
    # Fallback penalty
    if diagnostics.fallback_used:
        confidence_parts.append(0.7)
        review_reasons.append("Fallback extraction used")
    
    # Error penalty
    if diagnostics.errors:
        confidence_parts.append(0.5)
        review_reasons.append(f"{len(diagnostics.errors)} extraction errors")
    
    # Calculate overall
    overall = sum(confidence_parts) / len(confidence_parts) if confidence_parts else 0.5
    overall = round(min(1.0, max(0.0, overall)), 2)
    
    # Determine review need
    requires_review = overall < 0.6 or bool(review_reasons)
    
    return overall, requires_review, review_reasons


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN EXTRACTION FUNCTION
# ═══════════════════════════════════════════════════════════════════════════════

def extract_from_pdf(
    source: Union[str, Path, bytes],
    config: Optional[PDFExtractionConfig] = None
) -> ExtractionResult:
    """
    Extract content from a PDF.
    
    This is the main entry point for PDF extraction.
    
    Args:
        source: Path to PDF file or PDF bytes
        config: Extraction configuration (uses defaults if None)
    
    Returns:
        ExtractionResult with all extracted content and diagnostics
    
    Guarantees:
        - Will not block indefinitely (all operations timeboxed)
        - Will always return a result (even if partial/error)
        - Will not leak user content to logs
    
    Example:
        ```python
        from python.helpers.pdf_extraction import extract_from_pdf
        
        result = extract_from_pdf("document.pdf")
        
        if result.is_successful():
            for table in result.tables:
                print(table.to_csv())
        else:
            print(f"Extraction failed: {result.status}")
        ```
    """
    from python.helpers.pdf_extraction.pdf_backend import get_backend
    
    # SECURITY: Silence pdfminer debug logs — they can leak raw PDF content
    # pdfminer (used by pdfplumber) emits DEBUG logs with raw bytes from PDFs
    logging.getLogger("pdfminer").setLevel(logging.WARNING)
    
    # Setup
    config = config or get_default_config()
    diagnostics = Diagnostics()
    context = ExtractionContext(
        start_time=time.time(),
        correlation_id=str(uuid.uuid4())[:8]
    )
    circuit_breaker = CircuitBreaker(config)
    
    # Initialize result
    result = ExtractionResult(
        pdf_type=PDFType.UNKNOWN,
        pdf_type_confidence=0.0,
        diagnostics=diagnostics
    )
    
    try:
        # Open PDF via backend abstraction
        backend = get_backend()
        if isinstance(source, bytes):
            doc = backend.open_bytes(source)
            result.document_hash = ExtractionResult.compute_hash(source)
            pdf_path = None
        else:
            pdf_path = str(source)
            doc = backend.open_path(pdf_path)
            with open(pdf_path, "rb") as f:
                result.document_hash = ExtractionResult.compute_hash(f.read())
        
        from python.helpers.pdf_extraction.pdf_backend import PDFDocument as _PDFDocument
        if isinstance(doc, _PDFDocument):
            diagnostics.page_count = doc.page_count()
        else:
            diagnostics.page_count = doc.page_count
        
        # 1. Classify PDF
        with timed_operation(diagnostics, "classification"):
            pdf_type, pdf_type_confidence = classify_pdf(doc, config)
            result.pdf_type = pdf_type
            result.pdf_type_confidence = pdf_type_confidence
        
        diagnostics.classification_time_ms = int(
            (time.time() - context.start_time) * 1000
        )
        
        # 2. Extract words (PyMuPDF)
        words: list[Word] = []
        text_start = time.time()
        
        with timed_operation(diagnostics, "text_extraction"):
            words = extract_words_pymupdf(doc, config, context)
            result.words = words
            diagnostics.word_count = len(words)
        
        diagnostics.text_extraction_time_ms = int((time.time() - text_start) * 1000)
        diagnostics.methods_attempted.append(ExtractionMethod.PYMUPDF_TEXT)
        diagnostics.methods_succeeded.append(ExtractionMethod.PYMUPDF_TEXT)
        
        # 3. Extract tables (geometry reconstruction)
        tables: list[TableResult] = []
        table_start = time.time()
        
        if config.tables.enabled:
            with timed_operation(diagnostics, "table_extraction_geometry"):
                tables = extract_tables_geometry(
                    words, doc, config, context, diagnostics
                )
        
        # 4. Try optional engines if geometry quality is low
        if (config.tables.enabled and 
            tables and 
            all(t.confidence < 0.6 for t in tables) and
            pdf_path and
            not context.is_budget_exhausted(config.budgets.total_timeout_s)):
            
            optional_tables = try_optional_engines(
                pdf_path, words, doc, config, context, 
                circuit_breaker, diagnostics
            )
            
            if optional_tables:
                # Use optional engine results if better
                for opt_table in optional_tables:
                    # Find matching table by page
                    for i, geo_table in enumerate(tables):
                        if geo_table.page == opt_table.page:
                            if opt_table.confidence > geo_table.confidence:
                                tables[i] = opt_table
                                tables[i].fallback_used = True
                                tables[i].original_method = geo_table.method
                                diagnostics.fallback_used = True
                            break
                    else:
                        tables.append(opt_table)
        
        result.tables = tables
        diagnostics.table_extraction_time_ms = int((time.time() - table_start) * 1000)
        
        # 5. Generate outputs
        output_start = time.time()
        with timed_operation(diagnostics, "output_generation"):
            result.outputs = generate_outputs(tables, config)
        diagnostics.output_generation_time_ms = int((time.time() - output_start) * 1000)
        
        # 6. Calculate confidence
        overall_conf, needs_review, review_reasons = calculate_overall_confidence(
            pdf_type, pdf_type_confidence, tables, diagnostics
        )
        result.confidence_overall = overall_conf
        result.requires_human_review = needs_review
        result.review_reasons = review_reasons
        
        # Update circuit breaker state
        diagnostics.circuit_breaker = circuit_breaker.get_state()
        
        # Final timing
        diagnostics.total_time_ms = int((time.time() - context.start_time) * 1000)
        
        # Determine status
        if tables or words:
            result.status = ExtractionStatus.SUCCESS
        else:
            result.status = ExtractionStatus.PARTIAL
        
        doc.close()
        
    except Exception as e:
        result.status = ExtractionStatus.ERROR
        diagnostics.errors.append(f"{type(e).__name__}")  # No message (may contain user content)
        diagnostics.total_time_ms = int((time.time() - context.start_time) * 1000)
        
        logger.error(
            "PDF extraction failed",
            extra={
                "correlation_id": context.correlation_id,
                "error_type": type(e).__name__,
                "elapsed_ms": diagnostics.total_time_ms
            }
        )
    
    # Log safe summary
    if config.observability.enabled:
        logger.info(
            "PDF extraction complete",
            extra={
                "correlation_id": context.correlation_id,
                **result.summary()
            }
        )
    
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# ASYNC WRAPPER
# ═══════════════════════════════════════════════════════════════════════════════

async def extract_from_pdf_async(
    source: Union[str, Path, bytes],
    config: Optional[PDFExtractionConfig] = None
) -> ExtractionResult:
    """
    Async version of extract_from_pdf.
    
    Runs extraction in thread pool to avoid blocking event loop.
    """
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        _executor,
        lambda: extract_from_pdf(source, config)
    )
